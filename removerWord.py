from docx import Document


def remove_watermark_from_word(file_path):
    """Remove headers and footers from a Word document.

    The previous implementation attempted to call ``Paragraph.clear`` which is
    not available in ``python-docx`` and resulted in an ``AttributeError`` at
    runtime. This function now clears the text of each paragraph found in the
    header and footer sections and saves the result to a new file.
    """

    try:
        doc = Document(file_path)

        for section in doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                paragraph.text = ""

            footer = section.footer
            for paragraph in footer.paragraphs:
                paragraph.text = ""

        output_path = file_path.replace(".docx", "_no_watermark.docx")
        doc.save(output_path)
        return output_path
    except Exception as e:
        print(f"Error: {e}")
        return None
